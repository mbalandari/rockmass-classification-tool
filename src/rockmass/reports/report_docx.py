import os
from docx import Document
from docx.shared import Inches

from src.rockmass.plotting import (
    create_rmr_bar_chart,
    create_q_radar_chart,
    create_gsi_diagram,
    save_figure_as_png,
)


class DOCXReport:
    def __init__(self, result):
        # ClassificationResult object
        self.result = result

    # ---------------------------------------------------------
    # TABLE CREATION (DOCX version)
    # ---------------------------------------------------------
    def _create_table(self, doc, title, data_dict):
        """
        Creates a DOCX table for a breakdown dictionary.
        """
        doc.add_heading(title, level=2)

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parameter"
        hdr_cells[1].text = "Value"

        for key, value in data_dict.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

        doc.add_paragraph()  # spacing

    # ---------------------------------------------------------
    # MAIN REPORT GENERATION
    # ---------------------------------------------------------
    def generate(self, output_path: str):
        output_dir = os.path.dirname(output_path)
        os.makedirs(output_dir, exist_ok=True)

        doc = Document()
        doc.add_heading("Rock Mass Classification Report", level=1)

        # -----------------------------------------------------
        # SUMMARY
        # -----------------------------------------------------
        doc.add_paragraph(f"RMR: {self.result.rmr:.2f}")
        doc.add_paragraph(f"Q-System: {self.result.q_value:.2f}")
        doc.add_paragraph(f"GSI: {self.result.gsi:.2f}")
        doc.add_paragraph(
            f"Support Recommendations: {self.result.support_recommendations}"
        )

        doc.add_paragraph()

        # -----------------------------------------------------
        # TABLES
        # -----------------------------------------------------
        self._create_table(doc, "RMR Breakdown", self.result.rmr_breakdown)
        self._create_table(doc, "Q-System Breakdown", self.result.q_breakdown)
        self._create_table(doc, "GSI Breakdown", self.result.gsi_breakdown)

        # -----------------------------------------------------
        # GENERATE CHARTS
        # -----------------------------------------------------
        rmr_fig = create_rmr_bar_chart(self.result.rmr_breakdown)
        q_fig = create_q_radar_chart(self.result.q_breakdown)

        # IMPORTANT: pass actual GSI value
        gsi_fig = create_gsi_diagram(
            self.result.gsi_breakdown, gsi_value=self.result.gsi
        )

        rmr_path = os.path.join(output_dir, "rmr_chart.png")
        q_path = os.path.join(output_dir, "q_chart.png")
        gsi_path = os.path.join(output_dir, "gsi_chart.png")

        save_figure_as_png(rmr_fig, rmr_path)
        save_figure_as_png(q_fig, q_path)
        save_figure_as_png(gsi_fig, gsi_path)

        # -----------------------------------------------------
        # INSERT CHARTS
        # -----------------------------------------------------
        doc.add_page_break()
        doc.add_heading("RMR Breakdown Chart", level=2)
        doc.add_picture(rmr_path, width=Inches(6))

        doc.add_page_break()
        doc.add_heading("Q-System Chart", level=2)
        doc.add_picture(q_path, width=Inches(6))

        doc.add_page_break()
        doc.add_heading("GSI Diagram", level=2)
        doc.add_picture(gsi_path, width=Inches(6))

        # -----------------------------------------------------
        # SAVE DOCX
        # -----------------------------------------------------
        doc.save(output_path)
